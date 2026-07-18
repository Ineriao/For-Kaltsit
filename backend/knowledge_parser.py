import hashlib
import io
import json
import re
import threading
from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path


SUPPORTED_SOURCE_TYPES = {"txt", "md", "json", "html", "htm", "pdf", "docx", "png", "jpg", "jpeg", "webp", "bmp"}
IMAGE_SOURCE_TYPES = {"png", "jpg", "jpeg", "webp", "bmp"}


@dataclass(frozen=True)
class ParsedSection:
    content: str
    locator_text: str
    locator: dict


class KnowledgeParser:
    def __init__(self) -> None:
        self._ocr = None
        self._ocr_lock = threading.Lock()

    def parse(self, data: bytes, source_type: str) -> tuple[list[dict], dict]:
        normalized_type = source_type.lower().lstrip(".")
        if normalized_type not in SUPPORTED_SOURCE_TYPES:
            raise ValueError("不支持该资料格式")

        if normalized_type in {"txt", "md"}:
            sections = self._parse_text(self._decode_text(data))
        elif normalized_type == "json":
            sections = self._parse_json(self._decode_text(data))
        elif normalized_type in {"html", "htm"}:
            sections = self._parse_html(self._decode_text(data))
        elif normalized_type == "pdf":
            sections = self._parse_pdf(data)
        elif normalized_type == "docx":
            sections = self._parse_docx(data)
        else:
            sections = self._parse_image(data)

        chunks = self._chunk_sections(sections)
        metadata = {
            "bytes": len(data),
            "sections": len(sections),
            "ocr": normalized_type in IMAGE_SOURCE_TYPES or any(section.locator.get("ocr") for section in sections),
        }
        return chunks, metadata

    def ocr_status(self) -> dict:
        try:
            from rapidocr import RapidOCR

            return {"available": RapidOCR is not None, "engine": "RapidOCR"}
        except ImportError:
            return {"available": False, "engine": "RapidOCR"}

    @staticmethod
    def _decode_text(data: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-16", "gb18030"):
            try:
                return data.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise ValueError("无法识别文本文件编码")

    @staticmethod
    def _parse_text(content: str) -> list[ParsedSection]:
        lines = content.replace("\r\n", "\n").replace("\r", "\n").split("\n")
        sections = []
        start = 0
        buffer = []
        for index, line in enumerate([*lines, ""]):
            if line.strip():
                if not buffer:
                    start = index + 1
                buffer.append(line.strip())
                continue
            if not buffer:
                continue
            end = start + len(buffer) - 1
            locator_text = f"第 {start}–{end} 行" if end > start else f"第 {start} 行"
            sections.append(ParsedSection("\n".join(buffer), locator_text, {"line_start": start, "line_end": end}))
            buffer = []
        return sections

    @staticmethod
    def _parse_json(content: str) -> list[ParsedSection]:
        try:
            root = json.loads(content)
        except json.JSONDecodeError as error:
            raise ValueError("JSON 文件格式无效") from error

        sections = []

        def walk(value, path: str = "$") -> None:
            if isinstance(value, dict):
                for key, child in value.items():
                    walk(child, f"{path}.{key}")
            elif isinstance(value, list):
                for index, child in enumerate(value):
                    walk(child, f"{path}[{index}]")
            elif value is not None:
                sections.append(ParsedSection(str(value), path, {"json_path": path}))

        walk(root)
        return sections

    @staticmethod
    def _parse_html(content: str) -> list[ParsedSection]:
        parser = _ReadableHtmlParser()
        parser.feed(content)
        parser.close()
        return [
            ParsedSection(text, f"内容块 {index}", {"block": index})
            for index, text in enumerate(parser.blocks, 1)
        ]

    def _parse_pdf(self, data: bytes) -> list[ParsedSection]:
        try:
            from pypdf import PdfReader
        except ImportError as error:
            raise RuntimeError("PDF 解析组件未安装") from error

        try:
            reader = PdfReader(io.BytesIO(data))
        except Exception as error:
            raise ValueError("PDF 文件无法读取") from error

        sections = []
        for page_number, page in enumerate(reader.pages, 1):
            text = (page.extract_text() or "").strip()
            if text:
                sections.append(ParsedSection(text, f"第 {page_number} 页", {"page": page_number}))
                continue
            image_texts = []
            for image in getattr(page, "images", []):
                recognized = self._recognize_image(image.data)
                if recognized:
                    image_texts.append(recognized)
            if image_texts:
                sections.append(
                    ParsedSection("\n".join(image_texts), f"第 {page_number} 页 · OCR", {"page": page_number, "ocr": True})
                )
        return sections

    @staticmethod
    def _parse_docx(data: bytes) -> list[ParsedSection]:
        try:
            from docx import Document
        except ImportError as error:
            raise RuntimeError("DOCX 解析组件未安装") from error

        try:
            document = Document(io.BytesIO(data))
        except Exception as error:
            raise ValueError("DOCX 文件无法读取") from error

        sections = []
        for index, paragraph in enumerate(document.paragraphs, 1):
            text = paragraph.text.strip()
            if text:
                sections.append(ParsedSection(text, f"第 {index} 段", {"paragraph": index}))
        for table_index, table in enumerate(document.tables, 1):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                sections.append(ParsedSection("\n".join(rows), f"表格 {table_index}", {"table": table_index}))
        return sections

    def _parse_image(self, data: bytes) -> list[ParsedSection]:
        text = self._recognize_image(data)
        if not text:
            return []
        return [ParsedSection(text, "图片 OCR", {"ocr": True})]

    def _recognize_image(self, data: bytes) -> str:
        engine = self._load_ocr()
        try:
            result = engine(data)
        except Exception as error:
            raise ValueError("图片 OCR 识别失败") from error

        if hasattr(result, "txts"):
            return "\n".join(str(text).strip() for text in result.txts if str(text).strip())
        if isinstance(result, tuple):
            result = result[0]
        lines = []
        for item in result or []:
            if isinstance(item, (list, tuple)) and len(item) >= 2:
                text = item[1]
                if isinstance(text, (list, tuple)):
                    text = text[0]
                if str(text).strip():
                    lines.append(str(text).strip())
        return "\n".join(lines)

    def _load_ocr(self):
        if self._ocr is not None:
            return self._ocr
        with self._ocr_lock:
            if self._ocr is None:
                try:
                    from rapidocr import RapidOCR
                except ImportError as error:
                    raise RuntimeError("OCR 组件未安装") from error
                self._ocr = RapidOCR()
        return self._ocr

    @staticmethod
    def _chunk_sections(sections: list[ParsedSection], target_size: int = 420, overlap: int = 60) -> list[dict]:
        chunks = []
        for section in sections:
            normalized = re.sub(r"[ \t]+", " ", section.content)
            normalized = re.sub(r"\n{3,}", "\n\n", normalized).strip()
            if not normalized:
                continue
            step = max(1, target_size - overlap)
            for offset in range(0, len(normalized), step):
                content = normalized[offset:offset + target_size].strip()
                if len(content) < 2:
                    continue
                locator = dict(section.locator)
                if len(normalized) > target_size:
                    locator["character_start"] = offset + 1
                    locator["character_end"] = offset + len(content)
                chunks.append({
                    "content": content,
                    "locator_text": section.locator_text,
                    "locator": locator,
                    "content_hash": hashlib.sha256(content.encode("utf-8")).hexdigest(),
                })
        return chunks


class _ReadableHtmlParser(HTMLParser):
    BLOCK_TAGS = {"p", "div", "article", "section", "li", "blockquote", "h1", "h2", "h3", "h4", "h5", "h6", "tr"}
    IGNORED_TAGS = {"script", "style", "noscript", "svg"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.blocks = []
        self._buffer = []
        self._ignored_depth = 0

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in self.IGNORED_TAGS:
            self._ignored_depth += 1
        elif tag == "br" and not self._ignored_depth:
            self._buffer.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in self.IGNORED_TAGS:
            self._ignored_depth = max(0, self._ignored_depth - 1)
        elif tag in self.BLOCK_TAGS and not self._ignored_depth:
            self._flush()

    def handle_data(self, data: str) -> None:
        if not self._ignored_depth and data.strip():
            self._buffer.append(data.strip())

    def close(self) -> None:
        super().close()
        self._flush()

    def _flush(self) -> None:
        text = re.sub(r"[ \t]+", " ", " ".join(self._buffer)).strip()
        if text:
            self.blocks.append(text)
        self._buffer = []
